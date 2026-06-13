import streamlit as st
import pandas as pd
import json
import io
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- HELPER FUNCTIONS ---
def set_cell_shading(cell, fill):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def clean_housing(text):
    if not text or pd.isna(text): return ""
    return str(text).upper().replace("POD", "").strip()

def clean_webex(staff_name, staff_email, service_name):
    text = (str(staff_name) + " " + str(staff_email) + " " + str(service_name)).upper()
    mapping = {"ALPHA": "A", "APOD": "A", "BRAVO": "B", "BPOD": "B", "CHARLIE": "C", 
               "CPOD": "C", "WEBEXC": "C", "EDWARD": "E", "EPOD": "E", "EDWINA": "E", "FOXTROT": "F", 
               "FPOD": "F", "G POD": "G", "WEBEXG": "G", "INDIA": "I", "IPOD": "I", 
               "WEBEXI": "I", "CENTRAL": "CENTRAL"}
    for key, val in mapping.items():
        if key in text: return val
    import re
    match = re.search(r'([A-Z])[\s-]*(?:POD|WEBEX)', text)
    if match: return match.group(1)
    return "Check"

def parse_cf(cf_str):
    try: return json.loads(cf_str)
    except: return {}

# --- STREAMLIT WEB INTERFACE ---
st.set_page_config(page_title="Webex Scheduler", layout="wide")
st.title("📂 Webex Schedule Formatter")

uploaded_file = st.file_uploader("Choose a TSV file", type="tsv")

if uploaded_file is not None:
    df = pd.read_csv(uploaded_file, sep='\t')
    df['Name_Lower'] = df['Customer Name'].str.lower()
    
    # Process Merging
    processed_rows = []
    groups = df.groupby(['Date Time', 'Name_Lower', 'Customer Email'])

    for (dt_raw, name_low, email), group in groups:
        base = group.iloc[0]
        names, dobs = [], []
        for _, row in group.iterrows():
            cf = parse_cf(row[' Custom Fields'])
            names.extend([n.strip() for n in str(cf.get('INMATE NAME', '')).split(';') if n.strip()])
            dobs.extend([d.strip() for d in str(cf.get("Inmate's DOB", '')).split(';') if d.strip()])
        
        inmate_text = "\n".join([f"{names[i].upper()} (DOB: {dobs[i] if i < len(dobs) else 'N/A'})" 
                                for i in range(max(len(names), len(dobs)))])
        
        processed_rows.append({
            'time': datetime.strptime(dt_raw, '%m/%d/%Y %I:%M %p'),
            'inmates': inmate_text,
            'housing': clean_housing(parse_cf(base[' Custom Fields']).get('INMATE HOUSING LOCATION:', '')),
            'att_name': base['Customer Name'].upper(),
            'att_phone': base['Customer Phone'],
            'att_email': base['Customer Email'],
            'webex': clean_webex(base['Staff Name'], base['Staff Email'], base['Service'])
        })

    processed_rows.sort(key=lambda x: x['time'])
    
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    
    h_date = processed_rows[0]['time'].strftime('%m/%d/%Y')
    doc.add_heading(f'Webex Schedule {h_date}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    for i, h in enumerate(['Time', 'Inmate Information', 'Housing', 'Attorney Information', 'Webex']):
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold, run.font.size = True, Pt(11)

    # Shading variables
    last_time = None
    shading_on = False
    gray_hex = "E7E6E6"
    
    for r in processed_rows:
        if r['time'] != last_time:
            shading_on = not shading_on
            last_time = r['time']
            
        row_cells = table.add_row().cells
        if shading_on:
            for cell in row_cells: set_cell_shading(cell, gray_hex)
        
        row_cells[0].paragraphs[0].add_run((r['time'] + timedelta(minutes=15)).strftime('%I:%M %p')).font.size = Pt(11)
        row_cells[1].paragraphs[0].add_run(r['inmates']).font.size = Pt(11)
        row_cells[2].paragraphs[0].add_run(r['housing']).font.size = Pt(11)
        
        p_att = row_cells[3].paragraphs[0]
        p_att.add_run(f"{r['att_name']}\n").bold = True
        p_att.add_run(f"{r['att_phone']}\n").font.size = Pt(12)
        e_run = p_att.add_run(str(r['att_email']))
        e_run.font.size, e_run.font.name = Pt(12), 'Courier New'
        
        row_cells[4].paragraphs[0].add_run(r['webex']).font.size = Pt(11)

    bio = io.BytesIO()
    doc.save(bio)
    st.download_button("📥 Download Formatted Schedule", bio.getvalue(), f"Schedule_{h_date.replace('/','-')}.docx")
